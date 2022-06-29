import sys
from docx import Document

mes = input()
vr = input()
dev = [i[:-1] for i in sys.stdin]
doc = Document()
for i in dev:
    doc.add_heading(f"{i}, приглашаем тебя на праздник 8 марта", level=1) #заголовок
    mes1 = doc.add_paragraph("Праздник состоится ")#абзац текста этого док-та
    mes1.add_run(mes).bold = True#часть жирным шрифтом
    vr1 = doc.add_paragraph("Начало ")
    vr1.add_run(vr).bold = True
    doc.add_paragraph("Ждём тебя блин!")
    if i != dev[-1]:
        doc.add_page_break()
doc.save("Новый документ (6).docx")
