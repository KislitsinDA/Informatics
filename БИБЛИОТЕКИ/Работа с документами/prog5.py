from docx import Documents


def markdown_to_docs(text):
    text = text.split("\n")
    doc = Document()
    doc.add_heading(text[0], 0)
    for line in text[1:]:
        if not line:
            doc.add_paragraph("")
        if line.count("#") == 1:
            doc.add_heading(line[line.find(" "):], level=1)
        elif line.count("#") == 2:
            doc.add_heading(line[line.find(" "):], level=2)
        elif line.count("#") == 3:
            doc.add_heading(line[line.find(" "):], level=3)
        elif line.count("#") == 4:
            doc.add_heading(line[line.find(" "):], level=4)
        elif line.count("#") == 5:
            doc.add_heading(line[line.find(" "):], level=5)
        elif line.count("#") == 6:
            doc.add_heading(line[line.find(" "):], level=6)
        elif line[0] == "*" or line[0] == "+" or line[0] == "-":
            doc.add_paragraph(line[line.find(" "):], style="List Bullet")
        elif line[0].isdigit():
            doc.add_paragraph(line[line.find(" "):], style="List Number")
        elif line.count("*") == 2 or line.count("_") == 2:
            doc.add_paragraph(line[1:-1]).italic = True
        elif line.count("*") == 4 or line.count("_") == 4:
            doc.add_paragraph(line[1:-1]).bold = True
        elif line.count("*") == 6 or line.count("_") == 6:
            p = doc.add_paragraph(line[1:-1])
            p.italic = True
            p.bold = True
        else:
            doc.add_paragraph(line)
    doc.save("res.docx")
